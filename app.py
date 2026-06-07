import io
import os
import re
from typing import Any, Dict, List, Optional

import httpx
from docx import Document
from fastapi import FastAPI, Header, HTTPException, Query
from pypdf import PdfReader


app = FastAPI(
    title="Scripta Scientia OJS Bridge",
    version="1.4.0",
    description="Bridge seguro para conectar GPT Actions con OJS y extraer texto de manuscritos.",
)

OJS_BASE_URL = os.getenv("OJS_BASE_URL", "https://scriptascientia.com/sasc/api/v1").rstrip("/")
OJS_API_TOKEN = os.getenv("OJS_API_TOKEN")
BRIDGE_TOKEN = os.getenv("BRIDGE_TOKEN")


def check_auth(authorization: Optional[str]) -> None:
    if not BRIDGE_TOKEN:
        raise HTTPException(status_code=500, detail="BRIDGE_TOKEN not configured")

    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing Bearer token")

    token = authorization.replace("Bearer ", "", 1).strip()
    if token != BRIDGE_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid bridge token")


async def call_ojs(path: str, params: Optional[Dict[str, Any]] = None) -> Any:
    if not OJS_API_TOKEN:
        raise HTTPException(status_code=500, detail="OJS_API_TOKEN not configured")

    query = dict(params or {})
    query["apiToken"] = OJS_API_TOKEN

    url = f"{OJS_BASE_URL}{path}"

    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.get(url, params=query)

    if response.status_code >= 400:
        raise HTTPException(
            status_code=response.status_code,
            detail={
                "ojs_status_code": response.status_code,
                "ojs_response": response.text,
            },
        )

    try:
        return response.json()
    except ValueError:
        return {"raw": response.text}


async def download_file(url: str) -> bytes:
    if not OJS_API_TOKEN:
        raise HTTPException(status_code=500, detail="OJS_API_TOKEN not configured")

    params = {"apiToken": OJS_API_TOKEN}

    async with httpx.AsyncClient(timeout=60.0, follow_redirects=True) as client:
        response = await client.get(url, params=params)

    if response.status_code >= 400:
        raise HTTPException(
            status_code=response.status_code,
            detail={
                "download_status_code": response.status_code,
                "download_response": response.text[:1000],
            },
        )

    return response.content


def localized_value(value: Any) -> Any:
    if isinstance(value, dict):
        return value.get("es") or value.get("en") or next(iter(value.values()), None)
    return value


def get_items(data: Any) -> List[Dict[str, Any]]:
    if isinstance(data, dict) and isinstance(data.get("items"), list):
        return data["items"]
    if isinstance(data, list):
        return data
    return []


def clean_text(text: str) -> str:
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_texts = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    return clean_text("\n".join(paragraphs + table_texts))


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(io.BytesIO(content))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())
    return clean_text("\n\n".join(pages))


def word_count(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))


def detect_sections(text: str) -> Dict[str, bool]:
    lower = text.lower()

    patterns = {
        "title": r"\bt[ií]tulo\b",
        "abstract": r"\b(resumen|abstract)\b",
        "keywords": r"\b(palabras clave|keywords|descriptores)\b",
        "introduction": r"\b(introducci[oó]n|antecedentes)\b",
        "methods": r"\b(m[eé]todos|metodolog[ií]a|materiales y m[eé]todos)\b",
        "results": r"\b(resultados)\b",
        "discussion": r"\b(discusi[oó]n)\b",
        "conclusion": r"\b(conclusiones?|consideraciones finales)\b",
        "ethics": r"\b(comit[eé] de [eé]tica|aprobaci[oó]n [eé]tica|consentimiento informado)\b",
        "conflicts": r"\b(conflictos? de inter[eé]s|conflict of interest)\b",
        "funding": r"\b(financiamiento|fuente de financiamiento|funding)\b",
        "references": r"\b(referencias|bibliograf[ií]a|references)\b",
    }

    return {key: bool(re.search(pattern, lower)) for key, pattern in patterns.items()}


def summarize_submission(item: Dict[str, Any]) -> Dict[str, Any]:
    publications = item.get("publications") or []
    publication = publications[0] if publications else {}
    title = publication.get("title") or item.get("title") or item.get("fullTitle") or {}

    return {
        "id": item.get("id"),
        "status": item.get("status"),
        "statusLabel": item.get("statusLabel"),
        "dateSubmitted": item.get("dateSubmitted"),
        "lastModified": item.get("lastModified"),
        "title": localized_value(title),
        "currentPublicationId": item.get("currentPublicationId"),
        "stageId": item.get("stageId"),
        "url": item.get("_href"),
    }


def summarize_publication(publication: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": publication.get("id"),
        "title": localized_value(publication.get("title")),
        "abstract": localized_value(publication.get("abstract")),
        "keywords": publication.get("keywords"),
        "authorsString": publication.get("authorsString"),
        "datePublished": publication.get("datePublished"),
        "sectionId": publication.get("sectionId"),
        "doiObject": publication.get("doiObject"),
        "urlPath": publication.get("urlPath"),
    }


def summarize_participant(participant: Dict[str, Any]) -> Dict[str, Any]:
    user = participant.get("user") or participant

    return {
        "id": participant.get("id") or user.get("id"),
        "userId": participant.get("userId") or user.get("id"),
        "name": user.get("fullName") or user.get("name"),
        "email": user.get("email"),
        "roles": participant.get("roles") or participant.get("roleIds"),
        "stageAssignmentId": participant.get("stageAssignmentId"),
        "dateAssigned": participant.get("dateAssigned"),
    }


def summarize_file(file_item: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": file_item.get("id"),
        "submissionFileId": file_item.get("submissionFileId"),
        "name": localized_value(file_item.get("name")),
        "genreId": file_item.get("genreId"),
        "fileStage": file_item.get("fileStage"),
        "documentType": file_item.get("documentType"),
        "mimetype": file_item.get("mimetype"),
        "dateUploaded": file_item.get("dateUploaded"),
        "url": file_item.get("_href"),
    }


def choose_main_manuscript(files: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    if not files:
        return None

    def score(file_item: Dict[str, Any]) -> int:
        name = str(file_item.get("name") or "").lower()
        mimetype = str(file_item.get("mimetype") or "").lower()
        file_stage = file_item.get("fileStage")

        value = 0

        if ".docx" in name or "wordprocessingml" in mimetype:
            value += 50
        if ".doc" in name:
            value += 35
        if ".pdf" in name or "pdf" in mimetype:
            value += 25

        if any(term in name for term in ["manuscrito", "manuscript", "articulo", "artículo"]):
            value += 40

        if any(term in name for term in ["primera hoja", "hoja frontal", "cover", "title page"]):
            value -= 60

        if file_stage in [2, 4]:
            value += 10

        return value

    ranked = sorted(files, key=score, reverse=True)
    return ranked[0] if score(ranked[0]) > 0 else None


def detect_editorial_signals(
    submission_summary: Dict[str, Any],
    publications: List[Dict[str, Any]],
    participants: List[Dict[str, Any]],
    files: List[Dict[str, Any]],
) -> Dict[str, Any]:
    file_names = " | ".join(str(f.get("name") or "").lower() for f in files)

    return {
        "hasMainManuscript": any(
            term in file_names
            for term in ["articulo", "artículo", "manuscrito", "manuscript", ".docx", ".doc", ".pdf"]
        ),
        "hasFrontPageOrCoverFile": any(
            term in file_names
            for term in ["primera hoja", "hoja frontal", "front", "cover", "title page"]
        ),
        "hasAbstractInMetadata": any(bool(p.get("abstract")) for p in publications),
        "hasKeywordsInMetadata": any(bool(p.get("keywords")) for p in publications),
        "hasDoiInMetadata": any(bool(p.get("doiObject")) for p in publications),
        "fileCount": len(files),
        "participantCount": len(participants),
        "publicationVersionCount": len(publications),
        "currentStageId": submission_summary.get("stageId"),
        "currentStatus": submission_summary.get("status"),
    }


@app.get("/health")
def health():
    return {
        "status": "ok",
        "service": "scripta-ojs-bridge",
        "version": "1.4.0",
        "ojs_base_url": OJS_BASE_URL,
    }


@app.get("/issues")
async def list_issues(
    authorization: Optional[str] = Header(default=None),
    count: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
):
    check_auth(authorization)
    return await call_ojs("/issues", {"count": count, "offset": offset})


@app.get("/issues/{issue_id}")
async def get_issue(issue_id: int, authorization: Optional[str] = Header(default=None)):
    check_auth(authorization)
    return await call_ojs(f"/issues/{issue_id}")


@app.get("/submissions")
async def list_submissions(
    authorization: Optional[str] = Header(default=None),
    count: int = Query(default=5, ge=1, le=20),
    offset: int = Query(default=0, ge=0),
    searchPhrase: Optional[str] = Query(default=None),
    status: Optional[int] = Query(default=None),
):
    check_auth(authorization)

    params: Dict[str, Any] = {"count": count, "offset": offset}
    if searchPhrase:
        params["searchPhrase"] = searchPhrase
    if status is not None:
        params["status"] = status

    data = await call_ojs("/submissions", params)
    items = get_items(data)

    return {
        "items": [summarize_submission(item) for item in items],
        "itemsMax": data.get("itemsMax") if isinstance(data, dict) else None,
        "count": count,
        "offset": offset,
    }


@app.get("/submissions/{submission_id}")
async def get_submission(submission_id: int, authorization: Optional[str] = Header(default=None)):
    check_auth(authorization)
    return await call_ojs(f"/submissions/{submission_id}")


@app.get("/submissions/{submission_id}/publications")
async def list_submission_publications(
    submission_id: int,
    authorization: Optional[str] = Header(default=None),
):
    check_auth(authorization)
    return await call_ojs(f"/submissions/{submission_id}/publications")


@app.get("/submissions/{submission_id}/participants")
async def list_submission_participants(
    submission_id: int,
    authorization: Optional[str] = Header(default=None),
):
    check_auth(authorization)
    return await call_ojs(f"/submissions/{submission_id}/participants")


@app.get("/submissions/{submission_id}/files")
async def list_submission_files(
    submission_id: int,
    authorization: Optional[str] = Header(default=None),
    fileStage: Optional[int] = Query(default=None),
):
    check_auth(authorization)

    params: Dict[str, Any] = {}
    if fileStage is not None:
        params["fileStage"] = fileStage

    return await call_ojs(f"/submissions/{submission_id}/files", params)


@app.get("/editorial-review/{submission_id}")
async def editorial_review(submission_id: int, authorization: Optional[str] = Header(default=None)):
    check_auth(authorization)

    submission = await call_ojs(f"/submissions/{submission_id}")
    publications_raw = await call_ojs(f"/submissions/{submission_id}/publications")
    participants_raw = await call_ojs(f"/submissions/{submission_id}/participants")
    files_raw = await call_ojs(f"/submissions/{submission_id}/files")

    submission_summary = summarize_submission(submission if isinstance(submission, dict) else {})
    publications = [summarize_publication(p) for p in get_items(publications_raw)]
    participants = [summarize_participant(p) for p in get_items(participants_raw)]
    files = [summarize_file(f) for f in get_items(files_raw)]

    return {
        "submissionId": submission_id,
        "submission": submission_summary,
        "publications": publications,
        "participants": participants,
        "files": files,
        "editorialSignals": detect_editorial_signals(submission_summary, publications, participants, files),
    }


@app.get("/editorial-review-full/{submission_id}")
async def editorial_review_full(
    submission_id: int,
    authorization: Optional[str] = Header(default=None),
    maxCharacters: int = Query(default=30000, ge=1000, le=60000),
):
    check_auth(authorization)

    base_review = await editorial_review(submission_id, authorization)

    files = base_review.get("files", [])
    main_file = choose_main_manuscript(files)

    extraction = {
        "success": False,
        "file": main_file,
        "text": None,
        "wordCount": 0,
        "characterCount": 0,
        "sectionsDetected": {},
        "warning": None,
    }

    if not main_file:
        extraction["warning"] = "No se pudo identificar un manuscrito principal descargable."
        return {**base_review, "manuscriptExtraction": extraction}

    file_url = main_file.get("url")
    file_name = str(main_file.get("name") or "").lower()
    mimetype = str(main_file.get("mimetype") or "").lower()

    if not file_url:
        extraction["warning"] = "El archivo principal no incluye URL de descarga."
        return {**base_review, "manuscriptExtraction": extraction}

    content = await download_file(file_url)

    try:
        if ".docx" in file_name or "wordprocessingml" in mimetype:
            text = extract_docx_text(content)
        elif ".pdf" in file_name or "pdf" in mimetype:
            text = extract_pdf_text(content)
        else:
            extraction["warning"] = f"Tipo de archivo no soportado para extracción: {mimetype or file_name}"
            return {**base_review, "manuscriptExtraction": extraction}

        if len(text) > maxCharacters:
            text = text[:maxCharacters] + "\n\n[Texto truncado por límite de caracteres.]"

        extraction.update(
            {
                "success": True,
                "text": text,
                "wordCount": word_count(text),
                "characterCount": len(text),
                "sectionsDetected": detect_sections(text),
                "warning": None,
            }
        )

    except Exception as exc:
        extraction["warning"] = f"No se pudo extraer texto del manuscrito: {str(exc)}"

    return {**base_review, "manuscriptExtraction": extraction}


@app.get("/users")
async def list_users(
    authorization: Optional[str] = Header(default=None),
    count: int = Query(default=20, ge=1, le=100),
    offset: int = Query(default=0, ge=0),
    searchPhrase: Optional[str] = Query(default=None),
):
    check_auth(authorization)

    params: Dict[str, Any] = {"count": count, "offset": offset}
    if searchPhrase:
        params["searchPhrase"] = searchPhrase

    return await call_ojs("/users", params)


@app.get("/users/{user_id}")
async def get_user(user_id: int, authorization: Optional[str] = Header(default=None)):
    check_auth(authorization)
    return await call_ojs(f"/users/{user_id}")
